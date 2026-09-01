from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(__file__).resolve().parents[1] / 'docs' / 'Guia_Pruebas_Cliente_Potencial_EventPass.docx'
NAVY, GREEN, MINT, LIGHT, GRAY, RED = '073B4C', '009B77', 'E8F7F2', 'F3F6F7', '5B6770', 'B42318'

def shade(cell, color):
    tcpr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color); tcpr.append(shd)
def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); margins = tcPr.first_child_found_in('w:tcMar')
    if margins is None: margins = OxmlElement('w:tcMar'); tcPr.append(margins)
    for side, value in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = margins.find(qn(f'w:{side}'))
        if node is None: node = OxmlElement(f'w:{side}'); margins.append(node)
        node.set(qn('w:w'), str(value)); node.set(qn('w:type'), 'dxa')
def set_width(cell, twips):
    tcPr=cell._tc.get_or_add_tcPr(); width=tcPr.find(qn('w:tcW'))
    if width is None: width=OxmlElement('w:tcW'); tcPr.append(width)
    width.set(qn('w:w'),str(twips)); width.set(qn('w:type'),'dxa')
def set_font(run, size=11, bold=False, color='000000'):
    run.font.name='Calibri'; run._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); run._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri')
    run.font.size=Pt(size); run.bold=bold; run.font.color.rgb=RGBColor.from_string(color)
def add_text(p, text, **kw):
    r=p.add_run(text); set_font(r, **kw); return r
def add_link(p, text, url):
    part=p.part; rid=part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    h=OxmlElement('w:hyperlink'); h.set(qn('r:id'),rid); r=OxmlElement('w:r'); rPr=OxmlElement('w:rPr'); color=OxmlElement('w:color'); color.set(qn('w:val'),GREEN); rPr.append(color); u=OxmlElement('w:u'); u.set(qn('w:val'),'single'); rPr.append(u); r.append(rPr); t=OxmlElement('w:t'); t.text=text; r.append(t); h.append(r); p._p.append(h)
def style_para(p, before=0, after=6, line=1.15):
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=line
def heading(doc, text, level=1):
    p=doc.add_paragraph(style=f'Heading {level}'); style_para(p, before=16 if level==1 else 10, after=6)
    add_text(p,text, size=16 if level==1 else 13, bold=True, color=GREEN if level==1 else NAVY); return p
def body(doc, text, bold_lead=None):
    p=doc.add_paragraph(); style_para(p)
    if bold_lead and text.startswith(bold_lead): add_text(p,bold_lead,bold=True); add_text(p,text[len(bold_lead):])
    else: add_text(p,text)
    return p
def bullet(doc, text):
    p=doc.add_paragraph(style='List Bullet'); style_para(p, after=3); add_text(p,text); return p
def numbered(doc, text):
    p=doc.add_paragraph(style='List Number'); style_para(p, after=4); add_text(p,text); return p
def table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.style='Table Grid'; t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; shade(c,NAVY); set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths: set_width(c,widths[i])
        p=c.paragraphs[0]; style_para(p,after=0); add_text(p,h,size=9,bold=True,color='FFFFFF')
    for row in rows:
        cells=t.add_row().cells
        for i,value in enumerate(row):
            c=cells[i]; set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths: set_width(c,widths[i])
            p=c.paragraphs[0]; style_para(p,after=0); add_text(p,str(value),size=9.5)
    doc.add_paragraph().paragraph_format.space_after=Pt(3)
    return t
def callout(doc, title, text, color=MINT):
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.autofit=False; c=t.cell(0,0); shade(c,color); set_cell_margins(c,140,180,140,180); set_width(c,9360)
    p=c.paragraphs[0]; style_para(p,after=3); add_text(p,title,size=10,bold=True,color=NAVY)
    p=c.add_paragraph(); style_para(p,after=0); add_text(p,text,size=10)
    doc.add_paragraph().paragraph_format.space_after=Pt(3)
def page_break(doc): doc.add_page_break()

doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(0.7); sec.bottom_margin=Inches(0.7); sec.left_margin=Inches(0.75); sec.right_margin=Inches(0.75); sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35)
styles=doc.styles; styles['Normal'].font.name='Calibri'; styles['Normal']._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); styles['Normal'].font.size=Pt(11)
for s in ['List Bullet','List Number']:
    styles[s].font.name='Calibri'; styles[s]._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); styles[s].font.size=Pt(11)

header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; style_para(header,after=0); add_text(header,'EVENTPASS VE  |  GUÍA DE PRUEBAS PARA CLIENTE',size=8,bold=True,color=GRAY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; style_para(footer,after=0); add_text(footer,'Documento de evaluación - use datos de demostración.',size=8,color=GRAY)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(72); style_para(p,after=8); add_text(p,'EVENTPASS VE',size=13,bold=True,color=GREEN)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; style_para(p,after=10); add_text(p,'Guía de pruebas para cliente potencial',size=28,bold=True,color=NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; style_para(p,after=28); add_text(p,'Recorra el sistema como organizador de eventos. Marque lo que funciona y reporte cualquier detalle en la última sección.',size=12,color=GRAY)
table(doc,['Propósito','Tiempo sugerido','Qué necesita'],[['Conocer el sistema y evaluar si responde a su operación','45 a 60 minutos','Un computador, teléfono opcional y acceso de demostración']], [2600,2300,4460])
callout(doc,'Antes de comenzar','No necesita conocimientos técnicos. Use solo el evento y los datos preparados para la demostración. No comparta contraseñas y no elimine información existente.')
heading(doc,'Enlaces de la demostración',1)
links=[('Inicio','https://eventosfacil.net/'),('Acceso administrativo','https://eventosfacil.net/admin/login'),('Panel del organizador','https://eventosfacil.net/admin'),('Eventos','https://eventosfacil.net/admin/eventos'),('Acreditación','https://eventosfacil.net/admin/acreditacion'),('Check-in','https://eventosfacil.net/admin/checkin')]
for label,url in links:
    p=doc.add_paragraph(); style_para(p,after=3); add_text(p,label+': ',bold=True); add_link(p,url,url)
body(doc,'El administrador de EventPass entregará de forma privada el usuario de demostración y el nombre del evento preparado.')
page_break(doc)

heading(doc,'Cómo completar esta guía',1)
numbered(doc,'Siga los pasos de cada prueba en orden.')
numbered(doc,'Al terminar cada bloque, marque: Funciona, Necesita explicación o Presentó un error.')
numbered(doc,'Si aparece un error, no intente solucionarlo. Tome una captura y use la ficha de reporte al final.')
callout(doc,'Resultado esperado','La prueba se considera satisfactoria cuando la acción es clara, el resultado aparece sin esperas inesperadas y usted entiende qué hacer después.')
heading(doc,'Resumen de recorrido',2)
table(doc,['Bloque','Lo que comprobará','Resultado'],[
['1. Panel y evento','Acceso a las áreas de trabajo','□ Funciona  □ Revisar  □ Error'],
['2. Registro','Formulario, correo y credencial','□ Funciona  □ Revisar  □ Error'],
['3. Agenda pública','Horario, diseño y cambios visibles','□ Funciona  □ Revisar  □ Error'],
['4. Foros y exposición','Asientos, planos y espacios comerciales','□ Funciona  □ Revisar  □ Error'],
['5. Operación','Acreditación, impresión y check-in','□ Funciona  □ Revisar  □ Error'],
['6. Expositores y datos','Portal, visitantes, patrocinantes y reportes','□ Funciona  □ Revisar  □ Error'],
],[1650,4300,3410])

sections=[
('1. Panel del organizador y preparación del evento',[
('Entre al panel del organizador.', 'Se muestran tarjetas o accesos claros hacia Eventos, Registro, Acreditación, Check-in, Patrocinantes y equipo.'),
('Abra el evento preparado para la demo.', 'Puede encontrar las opciones de Agenda, Foro/Asientos, Expositores, Plano, Personal y configuración.'),
('Cambie un dato de prueba que el administrador haya indicado y guarde.', 'El mensaje confirma el guardado y el cambio permanece al volver a abrir.'),
]),
('2. Registro del participante, correo y credencial',[
('Abra el enlace público de registro del evento de demostración.', 'El formulario se entiende desde un teléfono o computador.'),
('Registre un participante con un correo de prueba autorizado.', 'Aparece una confirmación clara. Si es gratuito no habla de pago; si requiere pago, indica que debe cargar comprobante.'),
('Revise el correo recibido.', 'El contenido corresponde a la modalidad del evento y ofrece el siguiente paso correcto.'),
('Abra la credencial o QR del participante.', 'Se ve el nombre y un código utilizable para el ingreso.'),
]),
('3. Agenda y pantalla pública para el público',[
('Abra Agenda y cree o edite una actividad de demostración.', 'La actividad muestra hora, escenario, ponentes y patrocinante si corresponde.'),
('Entre a la pestaña Pantalla pública.', 'Puede publicar la agenda, definir título, colores, color del texto, fuente y tamaño.'),
('Abra la pantalla pública en una nueva pestaña.', 'Se ve como un tablero de programación: fecha, hora, actividad actual, próxima actividad y estados.'),
('Cambie el horario o cancele una actividad de demostración.', 'El cambio se refleja en la pantalla pública al actualizarse automáticamente.'),
]),
('4. Foro, asientos y reservas',[
('Abra Asientos/Foro.', 'Se ve el límite de aforo, ocupación y reservas institucionales si están configuradas.'),
('Cree una reserva de prueba para invitados o patrocinantes.', 'Las sillas quedan identificadas y se pueden liberar sin afectar las demás.'),
('Use una propuesta de plano con IA o el ejemplo Foro 120.', 'Se informa que se está construyendo; la propuesta aparece editable antes de aplicarla.'),
('Mueva una silla o ajuste una fila en el plano de demostración.', 'El plano sigue siendo coherente y los asientos se ajustan.'),
]),
('5. Exposición, stands y plano',[
('Abra Expositores y localice una empresa de prueba.', 'La empresa pertenece solo al evento abierto y se puede asociar a un stand.'),
('Abra el plano de exposición.', 'Los stands se entienden visualmente y se puede ver la empresa asignada.'),
('Si el administrador lo autorizó, cargue un plano de prueba PNG, JPG o PDF y use Analizar con IA.', 'La propuesta se revisa antes de aplicar. No aplique un resultado que no represente el plano.'),
('Abra el plano público si está publicado.', 'Un visitante puede identificar los stands y empresas.'),
]),
('6. Acreditación, impresión y check-in',[
('Abra Acreditación y busque el participante de demostración por nombre o correo.', 'El resultado aparece rápidamente y muestra una acción clara para acreditar o imprimir.'),
('Revise la configuración de impresión o realice una impresión de prueba si el equipo está listo.', 'El personal puede ver el flujo sin depender de conocimientos técnicos.'),
('Abra Check-in y registre la entrada por búsqueda manual o QR.', 'Se confirma el ingreso y una segunda lectura informa que la persona ya entró.'),
('Revise el contador de aforo.', 'Distingue registros, ingresos reales y capacidad disponible.'),
]),
('7. Portal del expositor, patrocinantes y visitantes',[
('Abra el portal de una empresa expositora.', 'La empresa ve su perfil, personal, tareas, actividades y visitantes del evento correcto.'),
('Cree o edite un integrante del personal de prueba.', 'Aparece en el listado y se puede editar, suspender o eliminar.'),
('Abra Escanear visitantes y lea una credencial QR del mismo evento.', 'Se registra una visita; una nueva visita posterior suma al contador sin mezclar eventos.'),
('Abra Patrocinantes y revise un acuerdo de demostración.', 'Se pueden identificar paquete, pagos, entregables y patrocinio de actividades.'),
]),
]
for title, steps in sections:
    page_break(doc); heading(doc,title,1)
    body(doc,'Siga cada acción. Al final del bloque marque el resultado global.')
    for i,(action,expected) in enumerate(steps,1):
        heading(doc,f'Paso {i}',2); body(doc,action); callout(doc,'Debe ocurrir',expected, LIGHT)
    table(doc,['Resultado del bloque','Comentario breve'],[['□ Funciona   □ Necesita explicación   □ Presentó un error','']], [3800,5560])

page_break(doc); heading(doc,'Reporte de errores y observaciones',1)
body(doc,'Use una ficha por cada situación. Si puede, adjunte una captura de pantalla o foto del teléfono. No escriba contraseñas, tokens, números de tarjetas ni datos sensibles de asistentes.')
for n in range(1,4):
    heading(doc,f'Incidencia {n}',2)
    table(doc,['Campo','Complete aquí'],[
        ['Fecha y hora',''],['Persona que realizó la prueba',''],['Evento y pantalla',''],['Qué estaba intentando hacer',''],['Qué esperaba que ocurriera',''],['Qué ocurrió realmente',''],['¿Impide continuar?','□ Sí   □ No'],['Captura o enlace adjunto',''],['Prioridad que percibe','□ Alta  □ Media  □ Baja'],
    ],[2800,6560])

page_break(doc); heading(doc,'Cierre de evaluación',1)
body(doc,'Gracias por dedicar tiempo a esta prueba. Sus observaciones ayudan a adaptar EventPass a la operación real de su organización.')
table(doc,['Pregunta','Respuesta'],[
['¿Qué le resultó más valioso?',''],['¿Qué función usaría primero en su próximo evento?',''],['¿Qué parte necesita que le expliquemos mejor?',''],['¿Qué información o integración faltaría para su operación?',''],['¿Desea una segunda sesión usando un evento real?','□ Sí   □ No   Fecha propuesta:'],
],[3600,5760])
callout(doc,'Entrega sugerida','Envíe este documento completo junto con las capturas al contacto de EventPass que coordinó la demostración.')

for section in doc.sections:
    section.page_width=Inches(8.5); section.page_height=Inches(11)
OUT.parent.mkdir(parents=True, exist_ok=True); doc.core_properties.title='Guía de pruebas para cliente potencial - EventPass VE'; doc.core_properties.author='EventPass VE'; doc.save(OUT)
print(OUT)
