from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Guia_Demo_y_Pruebas_EventPass_VE_ACTUALIZADA.docx"
ASSETS = ROOT / "docs" / "assets" / "demo"
GREEN = "059669"
INK = "18181B"
MUTED = "52525B"
LIGHT = "ECFDF5"

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd")) or OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    if shd.getparent() is None:
        tc_pr.append(shd)

def set_cell_text(cell, text, bold=False, color=INK):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def hyperlink(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), GREEN)
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single")
    rpr.extend([color, underline]); run.append(rpr)
    node = OxmlElement("w:t"); node.text = text; run.append(node); link.append(run)
    paragraph._p.append(link)

def add_callout(doc, title, body, fill=LIGHT):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.45)
    cell = table.cell(0, 0); shade(cell, fill)
    cell.margin_top = 120; cell.margin_bottom = 120
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title); r.bold = True; r.font.color.rgb = RGBColor.from_string(GREEN)
    p2 = cell.add_paragraph(body); p2.paragraph_format.space_after = Pt(0)
    return table

def add_steps(doc, steps):
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        p.add_run(step)

def add_check(doc, text="¿Funcionó como se esperaba?"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(f"{text}   ☐ Sí   ☐ No   ☐ Necesita ayuda")
    r.bold = True; r.font.color.rgb = RGBColor.from_string(MUTED)

def add_image(doc, filename, caption):
    path = ASSETS / filename
    if not path.exists():
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(6.35))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    c.runs[0].italic = True; c.runs[0].font.size = Pt(8); c.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(.65); section.bottom_margin = Inches(.65)
section.left_margin = Inches(.75); section.right_margin = Inches(.75)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"; normal.font.size = Pt(10); normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(5); normal.paragraph_format.line_spacing = 1.08
for name, size, color in [("Heading 1", 20, INK), ("Heading 2", 15, GREEN), ("Heading 3", 12, INK)]:
    st = styles[name]; st.font.name = "Aptos Display"; st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(10); st.paragraph_format.space_after = Pt(5); st.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.text = "EVENTPASS VE  ·  GUÍA DE DEMOSTRACIÓN Y PRUEBAS"
header.runs[0].font.size = Pt(8); header.runs[0].font.bold = True; header.runs[0].font.color.rgb = RGBColor.from_string(GREEN)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("Documento para partner comercial y equipo de eventos · Agosto 2026").font.size = Pt(8)

title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_before = Pt(14); title.paragraph_format.space_after = Pt(4)
r = title.add_run("Demo y pruebas de EventPass VE"); r.bold = True; r.font.name = "Aptos Display"; r.font.size = Pt(28); r.font.color.rgb = RGBColor.from_string(INK)
sub = doc.add_paragraph("Guía simple para presentar el sistema a un cliente y permitirle probarlo sin conocimientos técnicos.")
sub.runs[0].font.size = Pt(13); sub.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
add_callout(doc, "RECORRIDO RECOMENDADO · 25 A 30 MINUTOS", "Panel → Evento → Agenda y planos → Registro → Acreditación → Expositores → Portal del expositor → Visitantes del stand")
add_image(doc, "01-inicio-actual.png", "Inicio público de EventPass VE")

doc.add_heading("Antes de empezar", 1)
add_steps(doc, [
    "Recibe el usuario y la clave por un canal privado. No los escribas en este documento ni los muestres en pantalla.",
    "Usa la organización Asociación Demo y los eventos ya preparados. No elimines información existente.",
    "Abre cada enlace en una pestaña diferente para moverte con fluidez durante la presentación.",
    "Si algo falla, anota la pantalla, la acción, lo que esperabas y lo que ocurrió.",
])

doc.add_page_break()
doc.add_heading("Enlaces directos", 1)
links = [
    ("Acceso administrativo", "https://eventosfacil.net/admin/login"),
    ("Panel del organizador", "https://eventosfacil.net/admin"),
    ("Eventos", "https://eventosfacil.net/admin/eventos"),
    ("Expo Energia 2026", "https://eventosfacil.net/admin/eventos/276e4d25-b107-4393-9530-542db8ed03a3/administrar"),
    ("Registro público", "https://eventosfacil.net/e/276e4d25-b107-4393-9530-542db8ed03a3"),
    ("Plano público", "https://eventosfacil.net/expo/276e4d25-b107-4393-9530-542db8ed03a3/plano"),
    ("Expositores", "https://eventosfacil.net/admin/expositores/276e4d25-b107-4393-9530-542db8ed03a3"),
    ("Plano de exposición", "https://eventosfacil.net/admin/stands/276e4d25-b107-4393-9530-542db8ed03a3"),
    ("Foro con IA", "https://eventosfacil.net/admin/asientos/47ad0375-24dd-4f40-80c0-500f4362767c"),
    ("Acreditación", "https://eventosfacil.net/admin/acreditacion"),
    ("Check-in", "https://eventosfacil.net/admin/checkin"),
]
table = doc.add_table(rows=1, cols=2); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = "Table Grid"
set_cell_text(table.rows[0].cells[0], "Pantalla", True, "FFFFFF"); set_cell_text(table.rows[0].cells[1], "Enlace", True, "FFFFFF")
shade(table.rows[0].cells[0], GREEN); shade(table.rows[0].cells[1], GREEN)
for label, url in links:
    cells = table.add_row().cells; set_cell_text(cells[0], label, True)
    cells[1].text = ""; p = cells[1].paragraphs[0]; hyperlink(p, "Abrir pantalla", url)

doc.add_heading("Parte A · Presentación comercial", 1)
doc.add_heading("1. Panel principal", 2)
add_steps(doc, ["Inicia sesión y abre Panel.", "Explica las tarjetas: eventos, registros, acreditación, check-in, aliados y operación.", "Abre Eventos y vuelve al panel para mostrar que la navegación es directa."])
add_callout(doc, "FRASE SUGERIDA", "El organizador encuentra las partes principales de su operación en un solo lugar.")
add_image(doc, "06-dashboard.png", "Panel principal del organizador")
add_check(doc)

doc.add_heading("2. Preparar un evento", 2)
add_steps(doc, ["Abre Expo Energia 2026.", "Muestra nombre, fecha, sede y estado.", "Entra en Agenda y enseña una actividad.", "Abre Expositores y el Plano de exposición."])
add_image(doc, "07-evento.png", "Espacio de administración de un evento")
add_check(doc)

doc.add_heading("3. Registro del participante", 2)
add_steps(doc, ["Abre el registro público en otra pestaña.", "Explica que el participante completa sus datos desde el teléfono.", "Durante una demo no envíes un registro nuevo salvo que el administrador lo haya preparado.", "Vuelve al panel y muestra dónde aparecen los registros."])
add_image(doc, "04-registro-publico.png", "Formulario público de registro")
add_check(doc)

doc.add_heading("4. Acreditación y check-in", 2)
add_steps(doc, ["Abre Acreditación y selecciona Expo Energia 2026.", "Busca un participante preparado por nombre o correo.", "Abre la credencial y muestra la opción de impresión.", "Abre Check-in: el QR nuevo y los QR antiguos son compatibles.", "Repite el ingreso para comprobar que el sistema lo advierte."])
add_image(doc, "08-acreditacion.png", "Mostrador de acreditación")
add_check(doc)

doc.add_heading("5. Expositores y plano público", 2)
add_steps(doc, ["Abre Expositores y localiza una empresa con stand.", "Muestra el perfil enviado, su estado y el portal.", "Abre el plano público y selecciona una empresa del directorio.", "Explica que el visitante puede buscar empresas y ubicarlas."])
add_image(doc, "05-plano-publico.png", "Plano público y directorio de empresas")
add_check(doc)

doc.add_heading("6. Portal del expositor y visitantes", 2)
add_steps(doc, ["Desde Expositores, abre el portal de una empresa del evento.", "Entra en Escanear visitantes y selecciona el stand.", "Abre la cámara o usa el campo para lector USB.", "Escanea una credencial válida del mismo evento.", "Comprueba nombre, empresa, cargo, correo y contador.", "Repite después de 10 segundos y descarga el CSV."])
add_callout(doc, "RESULTADO ESPERADO", "Cada visita real queda registrada. Una lectura accidental inmediata no crea un duplicado.")
add_image(doc, "09-visitantes-stand.png", "Captación de visitantes desde el portal del expositor")
add_check(doc)

doc.add_heading("7. Planos con IA", 2)
add_steps(doc, ["Foro: abre el evento Foro, pulsa Foro 120 y Crear propuesta.", "Mientras se genera, comprueba que aparece el indicador de trabajo.", "Revisa escenario, capacidad, pasillo central y entradas antes de aplicar.", "Exposición: usa un plano vacío, carga PNG, JPG o PDF y pulsa Analizar con IA.", "Aplica una propuesta sólo cuando la correlación visual sea adecuada; después continúa editando manualmente."])
add_image(doc, "10-foro-ia.png", "Asistente de plano de foro con IA")
add_check(doc)

doc.add_heading("Parte B · Prueba autónoma del cliente", 1)
tests = [
    ("Acceso y navegación", "Entrar al panel; abrir Eventos, Acreditación, Check-in y volver al panel."),
    ("Evento y agenda", "Abrir Expo Energia 2026; localizar agenda, plano, expositores y personal."),
    ("Registro", "Abrir el formulario público y confirmar que los campos se entienden desde el teléfono."),
    ("Acreditación", "Buscar una persona, abrir su credencial y localizar impresión y configuración."),
    ("Plano público", "Buscar una empresa y ubicar su stand."),
    ("Portal expositor", "Abrir personal, pendientes, actividades y visitantes."),
    ("Visitantes del stand", "Escanear una credencial, repetir después de 10 segundos y descargar CSV."),
    ("Plano con IA", "Generar una propuesta de foro, revisarla y confirmar que sigue siendo editable."),
]
table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
for idx, text in enumerate(["Prueba", "Acción", "Resultado"]): set_cell_text(table.rows[0].cells[idx], text, True, "FFFFFF"); shade(table.rows[0].cells[idx], GREEN)
for name, action in tests:
    cells = table.add_row().cells; set_cell_text(cells[0], name, True); set_cell_text(cells[1], action); set_cell_text(cells[2], "☐ Bien\n☐ Revisar")

doc.add_heading("Cómo reportar una mejora", 2)
add_callout(doc, "FORMATO DE UNA LÍNEA", "En [pantalla], al intentar [acción], esperaba [resultado], pero ocurrió [resultado real].", "F4F4F5")
doc.add_heading("Cierre de la reunión", 2)
add_steps(doc, ["Pregunta qué parte del evento le quitaría más trabajo al cliente.", "Pregunta qué tipos de asistentes, expositores y patrocinantes maneja.", "Anota qué desea ver en una segunda sesión con datos reales.", "No prometas funciones que no se hayan mostrado."])

doc.core_properties.title = "Guía de demostración y pruebas - EventPass VE"
doc.core_properties.subject = "Recorrido comercial y pruebas manuales para usuarios de eventos"
doc.core_properties.author = "EventPass VE"
doc.save(OUT)
print(OUT)
