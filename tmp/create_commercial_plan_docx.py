from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Proyectos\Github\EventPass-VE")
SOURCE = ROOT / "docs" / "PLAN_COMERCIAL_EVENTO_ASOCIACION.md"
OUTPUT = ROOT / "docs" / "Plan_Comercial_EventPass_Asociacion.docx"

NAVY = "123047"
TEAL = "00866A"
TEAL_DARK = "006B55"
MINT = "E9F7F2"
GOLD = "D7A928"
GOLD_LIGHT = "FFF6D8"
INK = "18232B"
MUTED = "5F6B73"
LIGHT = "F3F6F7"
WHITE = "FFFFFF"
BORDER = "D5DEE2"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_font(run, size=None, bold=None, color=INK, italic=None, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = rgb(color)


def shade(element, fill):
    props = element.get_or_add_tcPr() if element.tag.endswith("tc") else element.get_or_add_pPr()
    shd = props.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        props.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            width = widths[idx]
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, 9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def add_rich_text(paragraph, text, base_size=10.5, base_color=INK):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`|\[.*?\]\(.*?\))", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, base_size, bold=True, color=base_color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, base_size - .3, color=TEAL_DARK, name="Consolas")
        else:
            link_match = re.fullmatch(r"\[(.*?)\]\((.*?)\)", part)
            if link_match:
                run = paragraph.add_run(link_match.group(1))
                set_font(run, base_size, color=TEAL_DARK)
                run.underline = True
            else:
                run = paragraph.add_run(part)
                set_font(run, base_size, color=base_color)


def add_callout(doc, label, value, note=None, fill=MINT, accent=TEAL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell._tc, fill)
    set_cell_margins(cell, top=190, bottom=190, start=240, end=240)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label.upper())
    set_font(r, 9, bold=True, color=accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(3 if note else 0)
    r2 = p2.add_run(value)
    set_font(r2, 23, bold=True, color=NAVY)
    if note:
        p3 = cell.add_paragraph()
        p3.paragraph_format.space_after = Pt(0)
        set_font(p3.add_run(note), 10, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, rows):
    columns = len(rows[0])
    if columns == 2:
        widths = [3600, 5760]
    elif columns == 3:
        widths = [2250, 3600, 3510]
    elif columns == 4:
        widths = [1700, 2700, 2400, 2560]
    else:
        widths = [9360 // columns] * columns
        widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=len(rows), cols=columns)
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    for ridx, row_data in enumerate(rows):
        row = table.rows[ridx]
        if ridx == 0:
            set_repeat_table_header(row)
        for cidx, value in enumerate(row_data):
            cell = row.cells[cidx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if ridx == 0:
                shade(cell._tc, NAVY)
            elif ridx % 2 == 0:
                shade(cell._tc, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            add_rich_text(p, value, 9.3, WHITE if ridx == 0 else INK)
            if ridx == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(.78)
    section.bottom_margin = Inches(.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(.35)
    section.footer_distance = Inches(.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for style_name, size, color, before, after in (
        ("Title", 31, NAVY, 0, 8),
        ("Subtitle", 14, MUTED, 0, 18),
        ("Heading 1", 17, NAVY, 16, 8),
        ("Heading 2", 13.5, TEAL_DARK, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display" if style_name != "Subtitle" else "Aptos"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.bold = style_name != "Subtitle"
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.left_indent = Inches(0)
        style.paragraph_format.right_indent = Inches(0)
        style.paragraph_format.first_line_indent = Inches(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(.375)
        style.paragraph_format.first_line_indent = Inches(-.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(p.add_run("EVENTPASS VE  |  PLAN COMERCIAL"), 8.5, bold=True, color=TEAL_DARK)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def build_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(25)
    p.paragraph_format.space_after = Pt(14)
    set_font(p.add_run("EVENTPASS VE"), 11, bold=True, color=TEAL)

    p = doc.add_paragraph(style="Title")
    p.add_run("Plan comercial")
    p = doc.add_paragraph(style="Subtitle")
    p.add_run("Evento gremial con exposición")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    set_font(p.add_run("Propuesta interna para estimar, presentar y negociar la primera contratación."), 11.5, color=MUTED)

    add_callout(doc, "Oferta de lanzamiento recomendada", "USD 1.200", "Precio de lista: USD 1.500 · Equipos y operación presencial se cotizan aparte")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run("CONTEXTO DEL CLIENTE"), 9, bold=True, color=TEAL_DARK)
    for text in (
        "Asociación gremial venezolana, no empresa organizadora de eventos.",
        "Participación gratuita para asistentes.",
        "Ingresos vinculados principalmente con la comercialización de stands.",
        "Cantidad de operadores, puntos de impresión y equipos todavía por confirmar.",
    ):
        p = doc.add_paragraph(style="List Bullet")
        add_rich_text(p, text)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run("Documento interno · 31 de agosto de 2026"), 9.5, color=MUTED)
    doc.add_page_break()


def parse_markdown(doc, text):
    lines = text.splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("## 1."))
    lines = lines[start:]
    i = 0
    in_code = False
    code_lines = []
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("```"):
            if in_code:
                table = doc.add_table(rows=1, cols=1)
                set_table_geometry(table, [9360])
                cell = table.cell(0, 0)
                shade(cell._tc, LIGHT)
                set_cell_margins(cell, 150, 180, 150, 180)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                set_font(p.add_run("\n".join(code_lines)), 9, color=NAVY, name="Consolas")
                doc.add_paragraph().paragraph_format.space_after = Pt(0)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        if line.startswith("## "):
            if line.startswith("## 3.") or line.startswith("## 8.") or line.startswith("## 12."):
                doc.add_page_break()
            doc.add_paragraph(line[3:], style="Heading 1")
            if line.startswith("## 12."):
                add_callout(
                    doc,
                    "Decisión recomendada",
                    "USD 1.200",
                    "Oferta fundadora bajo alcance controlado; completar dimensionamiento después de la reunión.",
                    fill=MINT,
                    accent=TEAL,
                )
        elif line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
        elif line.startswith("| "):
            table_lines = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            parsed = []
            for tline in table_lines:
                cells = [cell.strip() for cell in tline.strip("|").split("|")]
                if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    continue
                parsed.append(cells)
            add_table(doc, parsed)
            continue
        elif re.match(r"^- ", line):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, line[2:])
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("> "):
            table = doc.add_table(rows=1, cols=1)
            set_table_geometry(table, [9360])
            cell = table.cell(0, 0)
            shade(cell._tc, GOLD_LIGHT)
            set_cell_margins(cell, 160, 200, 160, 200)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_rich_text(p, line[2:].strip('"“”'), 11, NAVY)
            for run in p.runs:
                run.italic = True
            doc.add_paragraph().paragraph_format.space_after = Pt(0)
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            add_rich_text(p, line)
        else:
            p = doc.add_paragraph()
            add_rich_text(p, line)
        i += 1


def main():
    doc = Document()
    configure_document(doc)
    build_cover(doc)
    parse_markdown(doc, SOURCE.read_text(encoding="utf-8"))
    doc.core_properties.title = "Plan comercial - Evento gremial con exposición"
    doc.core_properties.subject = "Estimación comercial EventPass VE"
    doc.core_properties.author = "EventPass VE"
    doc.core_properties.keywords = "EventPass, evento, asociación, acreditación, stands"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
